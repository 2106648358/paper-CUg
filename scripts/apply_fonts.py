"""Post-process pandoc-generated docx to apply Chinese fonts"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document(r"D:\ProjectResources\paper-CUg\thesis.docx")

# Define font settings for each style
FONT_SETTINGS = {
    "Normal":           ("Times New Roman", "仿宋", Pt(12)),
    "Body Text":        ("Times New Roman", "仿宋", Pt(12)),
    "First Paragraph":  ("Times New Roman", "仿宋", Pt(12)),
    "Compact":          ("Times New Roman", "仿宋", Pt(12)),
    "Heading 1":        ("Times New Roman", "黑体", Pt(16)),
    "Heading 2":        ("Times New Roman", "黑体", Pt(14)),
    "Heading 3":        ("Times New Roman", "黑体", Pt(13)),
    "Heading 4":        ("Times New Roman", "黑体", Pt(12)),
    "Title":            ("Times New Roman", "黑体", Pt(22)),
    "Author":           ("Times New Roman", "仿宋", Pt(16)),
    "Date":             ("Times New Roman", "仿宋", Pt(12)),
    "Subtitle":         ("Times New Roman", "楷体", Pt(14)),
    "List Bullet":      ("Times New Roman", "仿宋", Pt(12)),
    "List Number":      ("Times New Roman", "仿宋", Pt(12)),
    "Block Text":       ("Times New Roman", "仿宋", Pt(12)),
}

def set_para_font(para, latin, ea, size):
    for run in para.runs:
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), latin)
        rFonts.set(qn('w:hAnsi'), latin)
        rFonts.set(qn('w:eastAsia'), ea)
        rFonts.set(qn('w:cs'), latin)
        sz = rpr.find(qn('w:sz'))
        if sz is None:
            sz = rpr.makeelement(qn('w:sz'), {})
            rpr.append(sz)
        sz.set(qn('w:val'), str(int(size.pt * 2)))
        szCs = rpr.find(qn('w:szCs'))
        if szCs is None:
            szCs = rpr.makeelement(qn('w:szCs'), {})
            rpr.append(szCs)
        szCs.set(qn('w:val'), str(int(size.pt * 2)))

def set_style_font(style, latin, ea, size, bold=False):
    """Set font at style level"""
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), latin)
    rFonts.set(qn('w:hAnsi'), latin)
    rFonts.set(qn('w:eastAsia'), ea)
    sz = rpr.find(qn('w:sz'))
    if sz is None:
        sz = rpr.makeelement(qn('w:sz'), {})
        rpr.append(sz)
    sz.set(qn('w:val'), str(int(size.pt * 2)))
    if bold:
        b = rpr.find(qn('w:b'))
        if b is None:
            b = rpr.makeelement(qn('w:b'), {})
            rpr.append(b)

# Apply font to each paragraph
total = len(doc.paragraphs)
processed = 0
for para in doc.paragraphs:
    style_name = para.style.name
    if style_name in FONT_SETTINGS:
        latin, ea, size = FONT_SETTINGS[style_name]
        # Set at style level first
        set_style_font(para.style, latin, ea, size)
        # Then apply to runs
        set_para_font(para, latin, ea, size)
        processed += 1

doc.save(r"D:\ProjectResources\paper-CUg\thesis.docx")
print(f"Done: {processed}/{total} paragraphs font-processed")
