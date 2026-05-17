"""Rebuild all tables in thesis_with_equations.docx as three-line tables"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import re

DOCX = r"D:\ProjectResources\paper-CUg\docx\thesis_with_equations.docx"
doc = Document(DOCX)

# ═══════════════════════════════════════
# TABLE DATA (from thesis.tex)
# ═══════════════════════════════════

TABLES = [
    # Table 2.1: 萨蒂1-9标度法
    {
        "search_heading": "2.2.2",
        "caption": "表2.1：萨蒂1-9标度法",
        "bookmark": "_tab_scale",
        "headers": ["标度", "含义"],
        "rows": [
            ["1", "同等重要"],
            ["3", "稍微重要"],
            ["5", "明显重要"],
            ["7", "强烈重要"],
            ["9", "极端重要"],
            ["2, 4, 6, 8", "中间值"],
        ],
    },
    # Table 2.2: 随机一致性指标RI
    {
        "search_heading": "2.2.3",
        "caption": "表2.2：随机一致性指标 RI",
        "bookmark": "_tab_ri",
        "headers": ["n", "3", "4", "5", "6", "7", "8", "9"],
        "rows": [
            ["RI", "0.52", "0.89", "1.12", "1.26", "1.36", "1.41", "1.46"],
        ],
    },
    # Table 3.1: 主要符号说明
    {
        "search_heading": "3.1",
        "caption": "表3.1：主要符号说明",
        "bookmark": "_tab_notation",
        "headers": ["符号", "含义"],
        "rows": [
            ["m", "评价对象数量，本研究 m = 12"],
            ["n", "评价指标数量，本研究 n = 6"],
            ["S = {s₁, …, sₘ}", "候选项目集合"],
            ["C = {c₁, …, cₙ}", "评价指标集合"],
            ["X = [xᵢⱼ]ₘ×ₙ", "决策矩阵，xᵢⱼ∈[0,1]"],
            ["vᵢ", "项目 sᵢ 的综合评分"],
            ["W = (w₁, …, wₙ)", "指标权重向量，Σwⱼ = 1"],
            ["Wᴬᴴᴾ", "AHP主观权重向量"],
            ["Wᴱᵂᴹ", "EWM客观权重向量"],
            ["α", "混合权重组合系数，α = 0.5 等权"],
            ["A", "AHP判断矩阵，aₚᵩ = 1/aₕₚ"],
            ["λₘₐₓ", "判断矩阵 A 的最大特征值"],
            ["CI, CR", "一致性指标与一致性比率"],
            ["x̃ᵢⱼ", "标准化决策矩阵元素"],
            ["pᵢⱼ", "比重矩阵元素"],
            ["Hⱼ", "指标 cⱼ 的信息熵，Hⱼ∈[0,1]"],
            ["Dⱼ", "信息效用值，Dⱼ = 1 − Hⱼ"],
            ["wⱼᴱᵂᴹ", "指标 cⱼ 的EWM客观权重"],
        ],
    },
    # Table 4.1: AHP主观权重
    {
        "search_heading": "4.2.1",
        "caption": "表4.1：AHP主观权重",
        "bookmark": "_tab_ahp",
        "headers": ["指标", "权重", "排名"],
        "rows": [
            ["流行度", "0.362", "1"],
            ["性别平等", "0.189", "2"],
            ["安全性", "0.189", "2"],
            ["可持续性", "0.106", "4"],
            ["包容性", "0.090", "5"],
            ["创新性", "0.064", "6"],
        ],
    },
    # Table 4.2: 六维评分决策矩阵 X
    {
        "search_heading": "4.2.2",
        "caption": "表4.2：12个代表性项目的六维评分决策矩阵 X（标准化前）",
        "bookmark": "_tab_decision",
        "headers": ["项目", "流行度", "性别平等", "可持续性", "包容性", "创新性", "安全性"],
        "rows": [
            ["田径", "0.814", "0.880", "0.955", "1.000", "0.027", "0.750"],
            ["游泳", "0.580", "0.950", "0.716", "0.432", "0.418", "0.880"],
            ["足球", "0.692", "0.800", "0.735", "0.756", "0.023", "0.500"],
            ["篮球", "0.503", "0.920", "0.626", "0.269", "0.985", "0.550"],
            ["攀岩", "0.349", "0.900", "0.568", "0.294", "0.985", "0.850"],
            ["滑板", "0.394", "0.900", "0.626", "0.294", "0.985", "0.350"],
            ["冲浪", "0.369", "0.900", "0.626", "0.269", "0.985", "0.800"],
            ["霹雳舞", "0.327", "0.920", "0.500", "0.110", "1.000", "0.700"],
            ["电子竞技", "0.950", "0.650", "0.850", "0.550", "1.000", "0.950"],
            ["板球", "0.205", "0.700", "0.500", "0.247", "0.023", "0.650"],
            ["空手道", "0.500", "0.900", "0.500", "0.185", "0.985", "0.600"],
            ["棒球/垒球", "0.445", "0.800", "0.590", "0.296", "0.377", "0.800"],
        ],
    },
    # Table 4.3: 标准化决策矩阵
    {
        "search_heading": "4.2.2",
        "caption": "表4.3：标准化决策矩阵 X̃（最小-最大归一化）",
        "bookmark": "_tab_normalized",
        "headers": ["项目", "流行度", "性别平等", "可持续性", "包容性", "创新性", "安全性"],
        "rows": [
            ["田径", "0.817", "0.767", "1.000", "1.000", "0.004", "0.667"],
            ["游泳", "0.503", "1.000", "0.475", "0.362", "0.404", "0.883"],
            ["足球", "0.654", "0.500", "0.516", "0.726", "0.000", "0.250"],
            ["篮球", "0.400", "0.900", "0.277", "0.179", "0.985", "0.333"],
            ["攀岩", "0.193", "0.833", "0.149", "0.207", "0.985", "0.833"],
            ["滑板", "0.254", "0.833", "0.277", "0.207", "0.985", "0.000"],
            ["冲浪", "0.220", "0.833", "0.277", "0.179", "0.985", "0.750"],
            ["霹雳舞", "0.164", "0.900", "0.000", "0.000", "1.000", "0.583"],
            ["电子竞技", "1.000", "0.000", "0.769", "0.494", "1.000", "1.000"],
            ["板球", "0.000", "0.167", "0.000", "0.154", "0.000", "0.500"],
            ["空手道", "0.396", "0.833", "0.000", "0.084", "0.985", "0.417"],
            ["棒球/垒球", "0.322", "0.500", "0.198", "0.209", "0.360", "0.750"],
        ],
    },
    # Table 4.4: 比重矩阵 P
    {
        "search_heading": "4.2.2",
        "caption": "表4.4：比重矩阵 P",
        "bookmark": "_tab_proportion",
        "headers": ["项目", "流行度", "性别平等", "可持续性", "包容性", "创新性", "安全性"],
        "rows": [
            ["田径", "0.166", "0.095", "0.254", "0.263", "0.001", "0.096"],
            ["游泳", "0.102", "0.124", "0.121", "0.095", "0.053", "0.127"],
            ["足球", "0.133", "0.062", "0.131", "0.191", "0.000", "0.036"],
            ["篮球", "0.081", "0.112", "0.070", "0.047", "0.128", "0.048"],
            ["攀岩", "0.039", "0.103", "0.038", "0.054", "0.128", "0.120"],
            ["滑板", "0.052", "0.103", "0.070", "0.054", "0.128", "0.000"],
            ["冲浪", "0.045", "0.103", "0.070", "0.047", "0.128", "0.108"],
            ["霹雳舞", "0.033", "0.112", "0.000", "0.000", "0.130", "0.084"],
            ["电子竞技", "0.203", "0.000", "0.195", "0.130", "0.130", "0.144"],
            ["板球", "0.000", "0.021", "0.000", "0.041", "0.000", "0.072"],
            ["空手道", "0.080", "0.103", "0.000", "0.022", "0.128", "0.060"],
            ["棒球/垒球", "0.065", "0.062", "0.050", "0.055", "0.047", "0.108"],
        ],
    },
    # Table 4.5: EWM客观权重
    {
        "search_heading": "4.2.2",
        "caption": "表4.5：EWM客观权重",
        "bookmark": "_tab_ewm",
        "headers": ["指标", "信息熵 Hⱼ", "效用值 Dⱼ", "权重 wⱼᴱᵂᴹ"],
        "rows": [
            ["流行度", "0.902", "0.098", "0.144"],
            ["性别平等", "0.940", "0.060", "0.088"],
            ["可持续性", "0.814", "0.186", "0.272"],
            ["包容性", "0.859", "0.141", "0.206"],
            ["创新性", "0.865", "0.135", "0.198"],
            ["安全性", "0.937", "0.063", "0.092"],
        ],
    },
    # Table 4.6: Shapiro-Wilk
    {
        "search_heading": "4.4",
        "caption": "表4.6：六维指标 Shapiro-Wilk 正态性检验结果",
        "bookmark": "_tab_shapiro",
        "headers": ["维度", "W 统计量", "p 值", "正态性 (α=0.05)"],
        "rows": [
            ["流行度", "0.941", "0.509", "通过"],
            ["性别平等", "0.817", "0.015", "未通过"],
            ["可持续性", "0.893", "0.128", "通过"],
            ["包容性", "0.825", "0.019", "未通过"],
            ["创新性", "0.722", "0.001", "未通过"],
            ["安全性", "0.971", "0.921", "通过"],
        ],
    },
    # Table 4.7: 权重扰动
    {
        "search_heading": "4.6.2",
        "caption": "表4.7：单维度权重 ±10% 扰动下排名变化",
        "bookmark": "_tab_perturbation",
        "headers": ["扰动维度", "方向", "受影响项目", "位移", "说明"],
        "rows": [
            ["流行度", "+10%", "滑板", "6→7", "与攀岩互换"],
            ["流行度", "−10%", "篮球", "5→4", "与足球互换"],
            ["可持续性", "+10%", "攀岩", "7→6", "可持续性略优"],
            ["包容性", "+10%", "攀岩", "7→6", "包容性原始分略高"],
            ["创新性", "−10%", "空手道", "8→9", "与冲浪互换"],
            ["安全性", "+10%", "篮球", "5→4", "安全性高于足球"],
        ],
    },
]

# ═══════════════════════════════════════
# PHASE 1: Clear old tables (keep algos + radar)
# ═══════════════════════════════════════

# Identify tables to keep: algorithm tables (contain "算法" in header) + radar table
tables_to_keep = []
for ti, table in enumerate(doc.tables):
    hdr = table.cell(0, 0).text.strip()
    if '算法' in hdr or '田径' in hdr:
        tables_to_keep.append(ti)

print(f"Keeping {len(tables_to_keep)} tables (algorithms + radar)")

# Remove all other tables (work backwards to avoid index shifts)
# First collect all table elements
all_table_elems = doc.element.body.findall(qn('w:tbl'))
for ti, tbl in enumerate(all_table_elems):
    if ti not in tables_to_keep:
        tbl.getparent().remove(tbl)

# Also remove old "Table Caption" paragraphs
for p in doc.paragraphs:
    if p.style.name == 'Table Caption' and p.text.strip():
        p._element.getparent().remove(p._element)

print("Old tables cleared")

# ═══════════════════════════════════════
# PHASE 2: Build paragraph position index
# ═══════════════════════════════════════

# Find insertion positions by heading search
def find_heading_pos(heading_text):
    """Find first paragraph after the matching heading"""
    found = False
    for i, p in enumerate(doc.paragraphs):
        if found and p.style.name.startswith('Heading'):
            return i  # next heading = end of section
        if p.style.name.startswith('Heading') and heading_text in p.text:
            found = True
            continue
        if found and p.style.name.startswith('Heading') == False and p.text.strip():
            return i  # first body paragraph after heading
    return None

# For multi-table sections (4.2.2), track position sequentially
last_pos = None
positions = {}
for tbl_data in TABLES:
    heading = tbl_data["search_heading"]
    # Find paragraph at end of its section
    found_heading = False
    next_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and heading in p.text:
            found_heading = True
            continue
        if found_heading:
            if p.style.name.startswith('Heading'):
                next_heading_idx = i
                break

    if next_heading_idx is None:
        # Last section, use end of doc
        next_heading_idx = len(doc.paragraphs) - 1

    # Use last position or 3 paragraphs before next heading
    if last_pos and heading == last_heading:
        pos = last_pos + 15  # insert 15 paras after last table
    else:
        pos = max(next_heading_idx - 3, 100)

    positions[heading] = pos
    last_pos = pos
    last_heading = heading

print(f"Found insertion positions for {len(positions)} sections")

# ═══════════════════════════════════════
# PHASE 3: Insert tables with three-line format
# ═══════════════════════════════════════

def add_bookmark(para_elem, name):
    bk_s = OxmlElement('w:bookmarkStart')
    bk_s.set(qn('w:id'), name)
    bk_s.set(qn('w:name'), name)
    bk_e = OxmlElement('w:bookmarkEnd')
    bk_e.set(qn('w:id'), name)
    para_elem.insert(0, bk_s)
    para_elem.append(bk_e)

def make_three_line(thick_inside=False):
    """Create three-line table borders XML"""
    t = f'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
    b = f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
    l = f'<w:left w:val="nil" w:sz="0" w:space="0" w:color="auto"/>'
    r = f'<w:right w:val="nil" w:sz="0" w:space="0" w:color="auto"/>'
    ih = f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    iv = f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    return f'<w:tblBorders {nsdecls("w")}>{t}{l}{b}{r}{ih}{iv}</w:tblBorders>'

def set_cell_border(cell, top=False, bottom=None, top_sz="12", bottom_sz="12"):
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove existing borders
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    parts = []
    parts.append(f'<w:left w:val="nil" w:sz="0" w:space="0" w:color="auto"/>')
    parts.append(f'<w:right w:val="nil" w:sz="0" w:space="0" w:color="auto"/>')
    if top:
        parts.append(f'<w:top w:val="single" w:sz="{top_sz}" w:space="0" w:color="000000"/>')
    if bottom:
        parts.append(f'<w:bottom w:val="single" w:sz="{bottom_sz}" w:space="0" w:color="000000"/>')
    tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}>{"".join(parts)}</w:tcBorders>'))

def insert_three_line_table(tbl_data, para_idx):
    """Create and insert a three-line table at paragraph position"""
    headers = tbl_data["headers"]
    rows = tbl_data["rows"]
    caption = tbl_data["caption"]
    bookmark = tbl_data["bookmark"]
    ncols = len(headers)
    nrows = len(rows) + 1  # +1 for header

    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove default table borders, set three-line
    tbl = table._tbl
    tblPr = tbl.get_or_add_tblPr()
    tblPr.append(parse_xml(make_three_line()))

    # Header row: bold, centered, 黑体 font
    for ci, hdr in enumerate(headers):
        cell = table.cell(0, ci)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'
        rpr = r._element.get_or_add_rPr()
        rpr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="黑体" w:ascii="Times New Roman"/>'))
        # Thick top + thin bottom on header
        set_cell_border(cell, top=True, bottom="4", top_sz="12", bottom_sz="4")

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'
            rpr = r._element.get_or_add_rPr()
            rpr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体" w:ascii="Times New Roman"/>'))
            # Remove side borders for all data cells
            tcP = cell._tc.get_or_add_tcPr()
            # First row after header also has thin top
            tcB = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="nil" w:sz="0"/><w:right w:val="nil" w:sz="0"/></w:tcBorders>')
            tcP.append(tcB)

    # Last row: thick bottom border
    for ci in range(ncols):
        set_cell_border(table.cell(nrows - 1, ci), bottom="12", bottom_sz="12")

    # Set column widths proportionally
    total_w = 14.5  # cm
    if ncols <= 3:
        col_w = total_w / ncols
        for ri in range(nrows):
            for ci in range(ncols):
                table.cell(ri, ci).width = Cm(col_w)
    else:
        # First col wider, rest equal
        first_w = 2.5
        rest_w = (total_w - first_w) / (ncols - 1)
        for ri in range(nrows):
            table.cell(ri, 0).width = Cm(first_w)
            for ci in range(1, ncols):
                table.cell(ri, ci).width = Cm(rest_w)

    # Create caption paragraph (centered, bold, 黑体)
    cap_para = OxmlElement('w:p')
    cap_pr = OxmlElement('w:pPr')
    cap_jc = OxmlElement('w:jc')
    cap_jc.set(qn('w:val'), 'center')
    cap_pr.append(cap_jc)
    cap_para.append(cap_pr)
    cap_run = OxmlElement('w:r')
    cap_para.append(cap_run)
    cap_rpr = cap_run.makeelement(qn('w:rPr'), {})
    cap_rpr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="黑体" w:ascii="Times New Roman"/>'))
    cap_b = cap_rpr.makeelement(qn('w:b'), {})
    cap_rpr.append(cap_b)
    cap_sz = cap_rpr.makeelement(qn('w:sz'), {})
    cap_sz.set(qn('w:val'), '21')
    cap_rpr.append(cap_sz)
    cap_run.append(cap_rpr)
    cap_t = OxmlElement('w:t')
    cap_t.text = caption
    cap_t.set(qn('xml:space'), 'preserve')
    cap_run.append(cap_t)
    add_bookmark(cap_para, bookmark)

    # Insert: table element first, then caption after it
    ref_para = doc.paragraphs[para_idx]._element
    ref_para.addnext(cap_para)
    ref_para.addnext(table._tbl)
    # Swap: we want TABLE then CAPTION, but currently CAPTION then TABLE
    # Remove and reorder
    parent = ref_para.getparent()
    parent.remove(table._tbl)
    parent.remove(cap_para)
    parent.insert(list(parent).index(ref_para) + 1, table._tbl)
    parent.insert(list(parent).index(ref_para) + 2, cap_para)

    return table

# Insert all tables
for tbl_data in TABLES:
    heading = tbl_data["search_heading"]
    pos = positions.get(heading, 200)
    insert_three_line_table(tbl_data, pos)
    print(f"  Inserted: {tbl_data['caption'][:60]}")

doc.save(DOCX)
print(f"\nDone: {len(TABLES)} tables rebuilt as three-line format")
