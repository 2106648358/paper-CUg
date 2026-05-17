"""Format algorithm pseudocode in thesis_with_equations.docx as three-line tables"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document(r"D:\ProjectResources\paper-CUg\docx\thesis_with_equations.docx")

algorithms = [
    {
        "title": "算法1\u2003AHP权重计算",
        "lines": [
            ("输入: 判断矩阵 A \u2208 \u211d\u207f\u02e3\u207f", False),
            ("输出: 权重向量 W\u1d40\u1d3c\u1d50, 一致性比率 CR", False),
            ("求解 AW = \u03bb\u2098\u2090\u2093 W", False),
            ("取 \u03bb\u2098\u2090\u2093 对应的特征向量 v", False),
            ("归一化 W\u1d40\u1d3c\u1d50 = v / \u03a3 v\u1d62", False),
            ("CI = (\u03bb\u2098\u2090\u2093\u2212n)/(n\u22121)", False),
            ("查表 RI, 计算 CR = CI / RI", False),
            ("if CR < 0.1 then", True),
            ("    return W\u1d40\u1d3c\u1d50, CR", True),
            ("else", True),
            ("    return 需调整判断矩阵", True),
            ("end if", True),
        ]
    },
    {
        "title": "算法2\u2003EWM权重计算",
        "lines": [
            ("输入: 决策矩阵 X \u2208 \u211d\u207f\u02e3\u1d50, 方向向量 d", False),
            ("输出: 权重向量 W\u1d31\u1d57\u1d39, 信息熵 H", False),
            ("for j = 1 to m do", True),
            ("    if d\u2c7c = 正向 then", True),
            ("        x\u0303\u1d62\u2c7c = (x\u1d62\u2c7c \u2212 min x)/(max x \u2212 min x)", True),
            ("    else", True),
            ("        x\u0303\u1d62\u2c7c = (max x \u2212 x\u1d62\u2c7c)/(max x \u2212 min x)", True),
            ("    end if", True),
            ("    p\u1d62\u2c7c = x\u0303\u1d62\u2c7c / \u03a3\u1d62 x\u0303\u1d62\u2c7c", True),
            ("    H\u2c7c = \u22121/ln n \u00b7 \u03a3\u1d62 p\u1d62\u2c7c ln p\u1d62\u2c7c", True),
            ("end for", True),
            ("D\u2c7c = 1 \u2212 H\u2c7c,  w\u2c7c\u1d31\u1d57\u1d39 = D\u2c7c / \u03a3 D\u2096", False),
            ("return W\u1d31\u1d57\u1d39, H", True),
        ]
    },
    {
        "title": "算法3\u2003混合权重与综合评分",
        "lines": [
            ("输入: W\u1d40\u1d3c\u1d50, W\u1d31\u1d57\u1d39, 组合系数 \u03b1, 决策矩阵 X", False),
            ("输出: 混合权重 W, 综合评分 V", False),
            ("W = \u03b1 \u00b7 W\u1d40\u1d3c\u1d50 + (1\u2212\u03b1) \u00b7 W\u1d31\u1d57\u1d39", False),
            ("W = W / \u03a3 W", False),
            ("for i = 1 to n do", True),
            ("    v\u1d62 = \u03a3\u2c7c w\u2c7c \u00b7 x\u0303\u1d62\u2c7c", True),
            ("end for", True),
            ("return W, V", True),
        ]
    },
]

# Find algorithm paragraph range - look for garbled text and surrounding context
algo_start = None
algo_end = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text
    # "需要调整判断矩阵" is the last line of Algorithm 1 garbled text
    if '需要调整判断矩阵' in txt and algo_start is None:
        algo_start = i
    if algo_start is not None:
        # "混合权重" and "综合评分" appear in Algorithm 3 garbled text
        if '混合权重' in txt and '综合评分' in txt:
            algo_end = i + 1
            break

# Fallback: if we didn't find exactly, try broader search
if algo_start is None:
    for i, p in enumerate(doc.paragraphs):
        if '判断矩阵' in p.text and '权值计算' in p.text and '一致性检验' in p.text:
            algo_start = i
            break

if algo_end is None:
    for i in range(algo_start+1, min(algo_start+10, len(doc.paragraphs))):
        if '混合权重' in doc.paragraphs[i].text:
            algo_end = i + 1
            break

print(f"Algorithm paragraphs: P{algo_start} to P{algo_end}")

if algo_start is None or algo_end is None:
    print("ERROR: searching broadly...")
    for i in range(240, 250):
        t = doc.paragraphs[i].text.strip()
        if t:
            print(f"  P{i}: {t[:100]}")
    exit(1)

# Collect paragraph elements to remove (before insertion)
old_elems = []
for i in range(algo_start, algo_end):
    old_elems.append(doc.paragraphs[algo_start]._element)

# Reference element: the first old paragraph (tables go before it)
ref_elem = old_elems[0]

def add_algorithm_table(algo, insert_before_elem):
    n_rows = len(algo["lines"]) + 1
    table = doc.add_table(rows=n_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Merge header
    table.cell(0, 0).merge(table.cell(0, 1))

    # Header
    hdr_para = table.cell(0, 0).paragraphs[0]
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_run = hdr_para.add_run(algo["title"])
    hdr_run.bold = True
    hdr_run.font.size = Pt(10.5)
    hdr_run.font.name = "Times New Roman"
    rpr = hdr_run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="黑体" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>')
    rpr.insert(0, rFonts)

    # Remove table-level borders
    tbl = table._tbl
    tblPr = tbl.get_or_add_tblPr()
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>')
    tblPr.append(tblBorders)

    def set_borders(cell, top=None, bottom=None):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        # also remove left/right
        for edge_name in ['left', 'right']:
            edge = parse_xml(f'<w:{edge_name} {nsdecls("w")} w:val="nil" w:sz="0" w:space="0" w:color="auto"/>')
            tcBorders.append(edge)
        if top:
            tcBorders.append(parse_xml(f'<w:top {nsdecls("w")} w:val="single" w:sz="12" w:space="0" w:color="000000"/>'))
        if bottom:
            tcBorders.append(parse_xml(f'<w:bottom {nsdecls("w")} w:val="single" w:sz="{bottom}" w:space="0" w:color="000000"/>'))
        tcPr.append(tcBorders)

    # Header row borders: thick top, thin bottom
    for cell in table.rows[0].cells:
        set_borders(cell, top=True, bottom="4")

    # Body rows
    for idx, (line_text, is_keyword) in enumerate(algo["lines"]):
        row_idx = idx + 1
        # Line number
        ln_para = table.cell(row_idx, 0).paragraphs[0]
        ln_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ln_run = ln_para.add_run(str(idx + 1))
        ln_run.font.size = Pt(10)
        ln_run.font.name = "Times New Roman"
        ln_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        if is_keyword:
            ln_run.bold = True

        # Code text
        cd_para = table.cell(row_idx, 1).paragraphs[0]
        cd_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cd_run = cd_para.add_run(line_text)
        cd_run.font.size = Pt(10)
        cd_run.font.name = "Times New Roman"
        cdrpr = cd_run._element.get_or_add_rPr()
        cdFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="宋体" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>')
        cdrpr.insert(0, cdFonts)
        if is_keyword:
            cd_run.bold = True

        # Remove left/right borders
        for cell in table.rows[row_idx].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            for en in ['left','right']:
                tcBorders.append(parse_xml(f'<w:{en} {nsdecls("w")} w:val="nil" w:sz="0" w:space="0" w:color="auto"/>'))
            tcPr.append(tcBorders)

    # Last row: thick bottom border
    for cell in table.rows[-1].cells:
        set_borders(cell, bottom="12")

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.2)
        row.cells[1].width = Cm(13.3)

    # Insert table before the reference element
    insert_before_elem.addprevious(table._tbl)

# Insert three algorithm tables (reverse order so they appear correctly)
for algo_data in reversed(algorithms):
    add_algorithm_table(algo_data, ref_elem)

# Remove old garbled paragraphs
for elem in old_elems:
    elem.getparent().remove(elem)

doc.save(r"D:\ProjectResources\paper-CUg\docx\thesis_with_equations.docx")
print("Done: 3 algorithms formatted as three-line tables")
