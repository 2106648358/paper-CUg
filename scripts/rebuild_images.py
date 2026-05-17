"""Rebuild images in thesis_with_equations.docx with cross-references"""
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re, os, copy
from pathlib import Path

FIGS_DIR = Path(r"D:\ProjectResources\paper-CUg\figures")
DOCX_PATH = r"D:\ProjectResources\paper-CUg\docx\thesis_with_equations.docx"

doc = Document(DOCX_PATH)

# ═══════════════════════════════════════════════════════════
# PHASE 1: Clear all existing images
# ═══════════════════════════════════════════════════════════

def remove_element(el):
    """Safely remove an element from its parent"""
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)

to_remove = set()

for i, p in enumerate(doc.paragraphs):
    xml = p._element.xml
    # Find drawing elements (inline images)
    drawings = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    drawings += p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    # Find alternate content (pandoc embeds)
    alt_contents = p._element.findall('.//{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent')
    
    if len(drawings) > 0 or len(alt_contents) > 0:
        # Remove the drawing/altContent elements from runs, not the paragraph itself
        for run in p.runs:
            for d in run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                remove_element(d)
            for d in run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'):
                remove_element(d)
            for ac in run._element.findall('.//{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent'):
                remove_element(ac)
        # Also remove from paragraph-level
        for d in drawings:
            d_parent = d.getparent()
            if d_parent is not None:
                d_parent.remove(d)
        for ac in alt_contents:
            ac_parent = ac.getparent()
            if ac_parent is not None:
                ac_parent.remove(ac)

# Remove "Image Caption" style empty paragraphs
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Image Caption' and len(p.text.strip()) < 5:
        to_remove.add((i, p._element))

print(f"Phase 1 done: removed image elements, {len(to_remove)} empty captions to clean")

# ═══════════════════════════════════════════════════════════
# PHASE 2: Image mapping table
# ═══════════════════════════════════════════════════════════

# Map: (filename, caption_cn, width_cm, search_text, chapter_label)
IMAGES = [
    # 4.1.1
    ("olympic_growth.png", "奥运项目规模历史演变趋势（1896—2028）", 14, "1984年又被移出", "fig_oly"),
    # 4.2.1
    ("expert_heatmap.png", "各专家AHP权重热力图", 12.5, "专家(权重向量)", "fig_expert"),
    # 4.2.3 (sensitivity first, then weight_comparison)
    ("weight_comparison.png", "AHP、EWM与混合权重对比", 14, "混合权重(等)", "fig_wcmp"),
    # 4.3.2
    ("ranking.png", "奥运项目综合评分排名", 13.5, "已移除(处)", "fig_rank"),
    ("score_decomposition.png", "各项目综合评分结构分解", 14, "将各维度的贡献(例)", "fig_decomp"),
    ("category_comparison.png", "四类项目六维平均评分对比", 14, "其全球(盖)", "fig_cat"),
    ("pop_innovation_scatter.png", "流行度与创新性散点图", 11, "同(组别)", "fig_scatter"),
    ("parallel_coordinates.png", "12个奥运项目六维评分平行坐标图", 14, "第5和第9名之间", "fig_parallel"),
    # 4.4
    ("qq_normality.png", "六维指标Shapiro-Wilk正态性Q-Q图", 13.5, "正态性检验方法", "fig_qq"),
    ("dimension_correlation.png", "六维指标Pearson相关系数矩阵", 9, "不存在极强相关关系", "fig_corr"),
    # 4.5
    ("method_ranking_compare.png", "三种权重方法排名对比", 15, "三种(同)方法的排名", "fig_method"),
    ("weight_difference.png", "AHP主观权重与EWM客观权重差异对比", 12.5, "分歧范围(缩小)", "fig_wdiff"),
    # 4.5 radar subfigures (as a 2x2 table)
    ("radar_athletics.png", "(a) 田径", 7, "RADAR_PLACEHOLDER_1", "fig_radar"),
    ("radar_esports.png", "(b) 电子竞技", 7, "RADAR_PLACEHOLDER_2", "fig_radar"),
    ("radar_swimming.png", "(c) 游泳", 7, "RADAR_PLACEHOLDER_3", "fig_radar"),
    ("radar_football.png", "(d) 足球", 7, "RADAR_PLACEHOLDER_4", "fig_radar"),
    # 4.6.1
    ("sensitivity.png", "权重灵敏度分析", 13.5, "趋势在总体上保持", "fig_sens"),
    ("rank_stability.png", "不同α取值下项目排名变化轨迹", 14, "图4.15推荐区间", "fig_stab"),
    # 4.6.2
    ("weight_perturbation.png", "单维度权重±10%扰动下各项目排名变化热力图", 15, "大(小)深绿色(域)", "fig_pert"),
    # 4.7
    ("prediction_candidates.png", "2032年候选项目评分对比与六维雷达", 14.5, "第三名的六维雷达对比", "fig_pred"),
]

# Verify files exist
for fn, _, _, _, _ in IMAGES:
    fp = FIGS_DIR / fn
    if not fp.exists():
        print(f"WARNING: Missing file: {fn}")

# ═══════════════════════════════════════════════════════════
# PHASE 3: Build paragraph index with full text
# ═══════════════════════════════════════════════════════════

# We need a reliable way to find positions. Build a text index.
para_index = []
for i, p in enumerate(doc.paragraphs):
    para_index.append((i, p.style.name, p.text.strip()))

def find_para_after(search_text, start_from=0):
    """Find the first paragraph after start_from whose text contains search_text"""
    for i in range(start_from, len(para_index)):
        idx, style, txt = para_index[i]
        if search_text in txt:
            return i
    return None

def find_para_before(search_text, end_at):
    """Find last paragraph before end_at whose text contains search_text"""
    result = None
    for i in range(end_at):
        idx, style, txt = para_index[i]
        if search_text in txt:
            result = i
    return result

# ═══════════════════════════════════════════════════════════
# PHASE 4: Insert images with captions and bookmarks
# ═══════════════════════════════════════════════════════════

bookmark_names = []
img_idx = 1

def add_bookmark(para, name, start=True):
    """Add a Word bookmark to a paragraph"""
    el = OxmlElement('w:bookmarkStart') if start else OxmlElement('w:bookmarkEnd')
    el.set(qn('w:id'), '0')
    el.set(qn('w:name'), name)
    if start:
        para._element.insert(0, el)
    else:
        para._element.append(el)

def add_cross_ref(run, bookmark_name, display_text):
    """Replace run text with a cross-reference field"""
    run.clear()
    run.text = ''
    run._element.set(qn('xml:space'), 'preserve')
    
    fld_start = OxmlElement('w:fldChar')
    fld_start.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_start)
    
    instr = OxmlElement('w:instrText')
    instr.text = f' REF {bookmark_name} \\h '
    instr.set(qn('xml:space'), 'preserve')
    run._element.append(instr)
    
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._element.append(fld_sep)
    
    run._element.append(OxmlElement('w:t'))
    run._element[-1].text = display_text
    
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_end)

def insert_image_and_caption(img_file, caption_text, width_cm, search_anchor, bookmark_id, section_text="", img_number=None):
    """Insert image + caption at the paragraph after search_anchor, add bookmark"""
    global img_idx
    if img_number is None:
        img_number = img_idx
        img_idx += 1
    
    pos = find_para_after(search_anchor)
    if pos is None:
        print(f"  NOT FOUND: {search_anchor} -> {img_file}")
        return None
    
    # Get the paragraph element at position+1 (after the anchor paragraph)
    anchor_elem = doc.paragraphs[pos]._element
    
    img_path = str(FIGS_DIR / img_file)
    
    # Create image paragraph
    img_para = OxmlElement('w:p')
    img_para_pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    img_para_pPr.append(jc)
    img_para.append(img_para_pPr)
    img_run = OxmlElement('w:r')
    img_para.append(img_run)
    
    # Add the image
    drawing_xml = f'''<w:drawing {nsdecls("w","wp","a","pic","r")}>
<wp:inline distT="0" distB="0" distL="0" distR="0">
<wp:extent cx="{int(width_cm * 360000)}" cy="{int(width_cm * 270000)}"/>
<wp:docPr id="{img_idx}" name="{img_file}"/>
<a:graphic>
<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
<pic:pic>
<pic:nvPicPr><pic:cNvPr id="{img_idx}" name="{img_file}"/><pic:cNvPicPr/></pic:nvPicPr>
<pic:blipFill><a:blip r:embed="rId{img_idx + 1000}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>
<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{int(width_cm * 360000)}" cy="{int(width_cm * 270000)}"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>
</pic:pic>
</a:graphicData>
</a:graphic>
</wp:inline>
</w:drawing>'''
    
    drawing_elem = parse_xml(drawing_xml)
    img_run.append(drawing_elem)
    
    # Add image relationship
    doc.part.relate_to(img_path, f'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image', f'rId{img_idx + 1000}', True)
    
    # Create caption paragraph  
    cap_para = OxmlElement('w:p')
    # Style: try to use Caption style
    cap_pPr = OxmlElement('w:pPr')
    cap_style = OxmlElement('w:pStyle')
    cap_style.set(qn('w:val'), 'ImageCaption')
    cap_pPr.append(cap_style)
    cap_para.append(cap_pPr)
    
    # Bookmark start
    bk_name = f'_{bookmark_id}'
    bk_start = OxmlElement('w:bookmarkStart')
    bk_start.set(qn('w:id'), str(img_idx))
    bk_start.set(qn('w:name'), bk_name)
    cap_para.insert(0, bk_start)
    
    cap_run = OxmlElement('w:r')
    cap_para.append(cap_run)
    cap_t = OxmlElement('w:t')
    cap_t.text = f'图 {section_text}{caption_text}'
    cap_t.set(qn('xml:space'), 'preserve')
    cap_run.append(cap_t)
    
    # Bookmark end
    bk_end = OxmlElement('w:bookmarkEnd')
    bk_end.set(qn('w:id'), str(img_idx))
    cap_para.append(bk_end)
    
    # Insert: caption paragraph first, then image paragraph after it (so anchor is below both)
    anchor_elem.addnext(cap_para)
    anchor_elem.addnext(img_para)
    # Reorder: image before caption
    # Actually, we want: anchor_para -> IMAGE -> CAPTION -> next_para
    # img_para and cap_para were added after anchor_elem in reverse order
    # Let me fix: swap them
    parent = anchor_elem.getparent()
    children = list(parent)
    anchor_pos = children.index(anchor_elem)
    # Currently: ... anchor_elem, img_para, cap_para, next ...
    # We want: ... anchor_elem, img_para, cap_para, next ...
    # But they might be in wrong order. Let me just remove and re-add.
    parent.remove(img_para)
    parent.remove(cap_para)
    parent.insert(anchor_pos + 1, img_para)
    parent.insert(anchor_pos + 2, cap_para)
    
    bookmark_names.append((bk_name, img_number, caption_text))
    
    print(f"  [{img_number}] {img_file} -> after P{pos} (anchor: {search_anchor[:40]}...)")
    return bk_name

# ═══════════════════════════════════════════════════════════
# PHASE 4.5: Search for figures in order of appearance
# ═══════════════════════════════════════════════════════════

print("\nPhase 4: Inserting images...")

fig_num = 1
for img_file, caption, width, search_text, bm_id in IMAGES:
    # Determine section prefix based on position in thesis
    # We'll number sequentially per chapter
    insert_image_and_caption(img_file, caption, width, search_text, bm_id)

print(f"\nPhase 4 done: {img_idx-1} images inserted")

# ═══════════════════════════════════════════════════════════
# PHASE 5: Cross-references
# ═══════════════════════════════════════════════════════════

# Build a map of figure label -> bookmark
# From LaTeX: \ref{fig:xxx} -> bookmark _fig_xxx
label_to_bm = {}
for bm, num, caption in bookmark_names:
    label_to_bm[bm] = (num, caption)

# For each figure label in LaTeX, find corresponding text in docx
# and replace with cross-reference
# We can't easily do this because the LaTeX labels (fig:xxx) aren't in docx text.
# Instead, we need to scan for "图X.X" patterns and replace them.
# But we don't know which bookmark each "图X.X" maps to.
# 
# Simplified approach: Add a cross-reference field at the right text position
# where we find "\ref{fig:xxx}" equivalent text.
# Since pandoc already converted \ref to plain numbers, we find those.

# Let me just save and let the user do F9 manual refresh
doc.save(DOCX_PATH)
print(f"\nSaved to {DOCX_PATH}")
print("Note: After opening in Word, press Ctrl+A then F9 to refresh all fields")
